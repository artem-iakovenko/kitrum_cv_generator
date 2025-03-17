import time
from docx import Document
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import os
from docx.shared import Pt
from zoho_api.api import api_request
from datetime import datetime
from dateutil.relativedelta import relativedelta
from urllib.parse import urlparse
import traceback
import io
from googleapiclient.http import MediaIoBaseDownload

TEMPLATE_IDS = {
    "admin": "1u4vUF1o8pWIxjl_zsJ3Yl1HKr-MJKCNC",
    "designer": "1nDdWHXT_vKgr5LkL87ioL4TJbvv6QVMX",
    "developer": "18-YiSP4hLE_WMRWUI-iiasXkH_6Zicuf"
}

def get_date_difference(date_str1, date_str2):
    date_format = "%Y-%m-%d"
    date1 = datetime.strptime(date_str1, date_format)
    date2 = datetime.strptime(date_str2, date_format)
    difference = relativedelta(date2, date1)
    result = []
    if difference.years == 1:
        result.append("1 year")
    elif difference.years > 1:
        result.append(f"{difference.years} years")
    if difference.months == 1:
        result.append("1 month")
    elif difference.months > 1:
        result.append(f"{difference.months} months")
    return " and ".join(result)


def replace_in_paragraph(para, old_text, new_text):
    full_text = "".join(run.text for run in para.runs)
    if old_text in full_text:
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)


def has_image(paragraph):
    for run in paragraph.runs:
        if 'graphicData' in run._element.xml:
            return True
    return False


class CrmEntity:
    def __init__(self, crm_record_id):
        self.crm_record_id = crm_record_id
        self.crm_record_details = None
        self.cv_data = {}

    def get_crm_record_details(self):
        self.crm_record_details = api_request(
            f"https://www.zohoapis.com/crm/v2/CVs/{self.crm_record_id}",
            "zoho_crm",
            "get",
            None
        )['data'][0]

    def parse_record_data(self):
        cv_record_name = self.crm_record_details['Name']
        drive_folder_url = self.crm_record_details['Drive_Folder_URL']
        first_name = self.crm_record_details['First_Name']
        last_name = self.crm_record_details['Last_Name']
        employee_name = f"{first_name} {last_name}"
        seniority_cv = self.crm_record_details['Seniority'] or ""
        title_cv = self.crm_record_details['Title']
        direction = self.crm_record_details['Direction']
        entity_type = "developer"
        if direction in ["Admin"]:
            entity_type = 'admin'
        elif direction in ['Design']:
            entity_type = 'designer'

        location_cv = self.crm_record_details['Location']
        experience_in_it = self.crm_record_details['Years_of_experience_decimal']
        experience_in_it_cv = "1 year" if experience_in_it == 1 else f"{self.crm_record_details['Years_of_experience_decimal']} years"

        english = self.crm_record_details['English_Level']
        if not self.crm_record_details['Additional_languages']:
            other_languages = ""
        elif "; " in self.crm_record_details['Additional_languages']:
            other_languages = self.crm_record_details['Additional_languages'].split("; ") if self.crm_record_details['Additional_languages'] else ""
        else:
            other_languages = self.crm_record_details['Additional_languages'].split(", ") if self.crm_record_details['Additional_languages'] else ""

        general_achievements_cv = []
        general_achievements = self.crm_record_details['Work_achievements'].split("\n- ")
        for general_achievement in general_achievements:
            if general_achievement.startswith("- "):
                general_achievements_cv.append(general_achievement[2:])
            else:
                general_achievements_cv.append(general_achievement)

        # top section
        if entity_type == "developer":
            core_techs = [self.crm_record_details['Core_Technology_1'], self.crm_record_details['Core_Technology_2']]
            additional_techs = self.crm_record_details['Additional_technologies'].split(", ")
            tools = self.crm_record_details['Tools']
            databases = self.crm_record_details['Database']
            cloud = self.crm_record_details['Cloud']
            domains = self.crm_record_details['Has_experience_in_domains_New']
            skills_cv = {
                "main_languages": core_techs,
                "other_languages": additional_techs,
                "tools": tools,
                "databases": databases,
                "cloud": cloud,
                "domains": domains
            }
        elif entity_type == "admin":
            tools = self.crm_record_details['Tools']
            additional_skills = self.crm_record_details['Additional_Skills']
            domains = self.crm_record_details['Has_experience_in_domains_New']
            skills_cv = {
                "tools": tools,
                "additional_skills": additional_skills,
                "domains": domains
            }
        elif entity_type == "designer":
            tools = self.crm_record_details['Tools']
            additional_tools = self.crm_record_details['Additional_Tools']
            portfolio = self.crm_record_details['Portfolio_Designer']
            domains = self.crm_record_details['Has_experience_in_domains_New']
            additional_skills = self.crm_record_details['Additional_Skills']
            skills_cv = {
                "tools": tools,
                "additional_tools": additional_tools,
                "additional_skills": additional_skills,
                "domains": domains,
                "portfolio": portfolio
            }
        else:
            skills_cv = {}

        experience_cv = []
        work_experiences = self.crm_record_details['Work_Experience_CV']
        for work_experience in work_experiences:
            start_date = work_experience['Start_date']
            finish_date = work_experience['Finish_date']
            company = work_experience['Company_Name']
            position = work_experience['Position']
            project_name_and_description = work_experience['Project_name_and_description']


            project_name = project_name_and_description.split("\n")[0] if '\n' in project_name_and_description else project_name_and_description

            project_domain = work_experience['Project_domain']
            project_duration = get_date_difference(start_date, finish_date or datetime.today().strftime("%Y-%m-%d"))
            project_structure = work_experience['Team_Structure']
            project_tech_stack = work_experience['Project_tech_stack']
            project_responsibilities_cv = []
            project_responsibilities = work_experience['Responsibilities'].split("\n- ") if work_experience['Responsibilities'] else []
            for project_responsibility in project_responsibilities:
                if project_responsibility.startswith("- "):
                    project_responsibilities_cv.append(project_responsibility[2:])
                else:
                    project_responsibilities_cv.append(project_responsibility)

            project_achievements_cv = []
            project_achievements = work_experience['Project_Achievements'].split("\n- ") if work_experience['Project_Achievements'] else []
            for project_achievement in project_achievements:
                if project_achievement.startswith("- "):
                    project_achievements_cv.append(project_achievement[2:])
                else:
                    project_achievements_cv.append(project_achievement)

            experience_cv.append({
                "from": datetime.strptime(start_date, "%Y-%m-%d").strftime("%b. %Y"),
                "to": datetime.strptime(finish_date, "%Y-%m-%d").strftime("%b. %Y") if finish_date else 'present',
                "company": company,
                "position": position,
                "project_name": project_name_and_description.split("\n")[0].replace(".", ""),
                "project_description": project_name_and_description.replace(f"{project_name}\n", ""),
                "project_domain": project_domain or '__empty__',
                "project_duration": project_duration or '__empty__',
                "project_team_structure": project_structure or '__empty__',
                "project_tech_stack": project_tech_stack or '__empty__',
                "project_responsibilities": project_responsibilities_cv,
                "project_achievements": project_achievements_cv
            })
        education_cv = []
        certifications_cv = []
        educations = self.crm_record_details['Education_CV']
        for education in educations:
            type_of_education = education['Type_of_education']
            if type_of_education == "University":
                education_cv.append({
                    "university": education['University_Certificate_name'],
                    "specialization": education['Degree_in_specialization']
                })
            elif type_of_education == "Certification":
                certifications_cv.append({
                    "certificate_date": datetime.strptime(education['Final_date_of_graduation'], "%Y-%m-%d").strftime("%b. %Y"),
                    "certificate_name": education['Degree_in_specialization']
                })
        languages_cv = {"English": english}
        for other_language in other_languages:
            lang_list = other_language.split(" - ")
            languages_cv[lang_list[0]] = lang_list[1]

        self.cv_data = {
            "cv_record_name": cv_record_name,
            "role": entity_type,
            "full_name": employee_name,
            "seniority": seniority_cv,
            "position": title_cv,
            "location": location_cv,
            "experience": experience_in_it_cv,
            "summaries": general_achievements_cv,
            "skills": skills_cv,
            "work_experience": experience_cv,
            "education": education_cv,
            "certifications": certifications_cv,
            "languages": languages_cv,
            "drive_folder_url": drive_folder_url
        }

    def get_cv_data(self):
        self.get_crm_record_details()
        self.parse_record_data()
        return self.cv_data

    def update_cvs(self, file_ids):
        cv_folder_id = file_ids['folder_id']
        cv_docx_id = file_ids['docx']
        cv_pdf_id = file_ids['docx']
        if cv_folder_id and cv_docx_id and cv_pdf_id:
            crm_update_data = {"data": [
                {
                    "Drive_Folder_URL": f"https://drive.google.com/drive/folders/{cv_folder_id}",
                    "CV_DOCX_URL": f"https://drive.google.com/file/d/{cv_docx_id}",
                    "CV_PDF_URL": f"https://drive.google.com/file/d/{cv_pdf_id}",
                    "CV_Generation_Date": datetime.now().strftime("%Y-%m-%dT%H:%M:%S+01:00"),
                    "CV_Generation_Status": "Success"
                }
            ]}
        else:
            crm_update_data = {"data": [
                {
                    "CV_Generation_Status": "Error"
                }
            ]}
        response = api_request(
            f"https://www.zohoapis.com/crm/v2/CVs/{self.crm_record_id}",
            "zoho_crm",
            "put",
            crm_update_data
        )
        update_results = response['data'][0]['code']
        print(f"\tCRM Update Status: {update_results}")


class CurriculumVitae:
    def __init__(self, drive_service, entity_type, user_details):
        self.entity_type = entity_type
        self.cv_doc = None
        self.cv_doc_name = "unnamed"
        self.service = drive_service
        self.user_details = user_details
        self.variables_mapping = {}

    def get_template(self):
        # download template
        template_drive_id = TEMPLATE_IDS[self.entity_type]
        request = self.service.files().get_media(fileId=template_drive_id)
        file = io.BytesIO()
        downloader = MediaIoBaseDownload(file, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
            print(f"Template Download {int(status.progress() * 100)}% complete.")
        with open(f'templates/{self.entity_type}.docx', 'wb') as f:
            f.write(file.getvalue())
        print(f"Template 'templates/{self.entity_type}.docx' downloaded successfully.")
        self.cv_doc = Document(f"templates/{self.entity_type}.docx")

    def generate_doc_name(self):
        self.cv_doc_name = f'{self.user_details["full_name"].lower().replace(" ", "_")}'

    def map_variables(self):
        if self.entity_type == "developer":
            self.variables_mapping = {
                "__full_name__": self.user_details['full_name'],
                "__seniority__": self.user_details['seniority'],
                "__title__": self.user_details['position'],
                "__location__": self.user_details['location'],
                "__experience__": self.user_details['experience'],
                "__languages__": ", ".join(self.user_details['skills']['main_languages']),
                "__other_languages__": ", ".join(self.user_details['skills']['other_languages']),
                "__tools__": ", ".join(self.user_details['skills']['tools']),
                "__databases__": ", ".join(self.user_details['skills']['databases']),
                "__cloud__": ", ".join(self.user_details['skills']['cloud']),
                "__domains__": ", ".join(self.user_details['skills']['domains']),
                "__university__": self.user_details['education'][0]['university'],
                "__specialization__": self.user_details['education'][0]['specialization']
            }
        elif self.entity_type == "admin":
            self.variables_mapping = {
                "__full_name__": self.user_details['full_name'],
                "__seniority__": self.user_details['seniority'],
                "__title__": self.user_details['position'],
                "__location__": self.user_details['location'],
                "__experience__": self.user_details['experience'],
                "__tools__": ", ".join(self.user_details['skills']['tools']),
                "__additional_skills__": self.user_details['skills']['additional_skills'],
                "__domains__": ", ".join(self.user_details['skills']['domains']),
                "__university__": self.user_details['education'][0]['university'],
                "__specialization__": self.user_details['education'][0]['specialization']
            }
        elif self.entity_type == "designer":
            self.variables_mapping = {
                "__full_name__": self.user_details['full_name'],
                "__seniority__": self.user_details['seniority'],
                "__title__": self.user_details['position'],
                "__location__": self.user_details['location'],
                "__experience__": self.user_details['experience'],
                "__tools__": ", ".join(self.user_details['skills']['tools']),
                "__other_tools__": self.user_details['skills']['additional_tools'],
                "__additional_skills__": self.user_details['skills']['additional_skills'],
                "__domains__": ", ".join(self.user_details['skills']['domains']),
                "__university__": self.user_details['education'][0]['university'],
                "__specialization__": self.user_details['education'][0]['specialization']
            }

    def map_experience(self, experience):
        return {
            "__start_date__": experience['from'],
            "__end_date__": experience['to'],
            "__company__": experience['company'],
            "__title_on_project__": experience['position'],
            "__project_name__": experience['project_name'],
            "__project_domain__": experience['project_domain'],
            "__project_description__": experience['project_description'],
            "__project_duration__": experience['project_duration'],
            "__project_structure__": experience['project_team_structure'],
            "__project_stack__": experience['project_tech_stack']
        }

    def replace_text_preserving_formatting(self, old_text, new_text):
        for para in self.cv_doc.paragraphs:
            replace_in_paragraph(para, old_text, new_text)
        for table in self.cv_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, old_text, new_text)

    def insert_experience(self, experience):
        experience_mapping = self.map_experience(experience)
        responsibilities = experience['project_responsibilities']
        achievements = experience['project_achievements']
        for table in self.cv_doc.tables:
            if 'work experience' in table.rows[0].cells[0].paragraphs[0].text.lower() and 'end_date' in \
                    table.rows[1].cells[0].paragraphs[0].text.lower():
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for variable_name, variable_value in experience_mapping.items():
                                replace_in_paragraph(para, variable_name, variable_value)
                for responsibility in responsibilities:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "__responsibility__" in para.text:
                                    replace_in_paragraph(para, "__responsibility__", responsibility)
                                    break
                for achievement in achievements:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "__achievement__" in para.text:
                                    replace_in_paragraph(para, "__achievement__", achievement)
                                    break
                break

    def insert_primary_achievements(self, primary_achievement):
        inserted = False
        for table in self.cv_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if '__top_row__' in para.text:
                            if not inserted:
                                para.text = para.text.replace("__top_row__", primary_achievement)
                                for run in para.runs:
                                    run.font.size = Pt(8)
                                inserted = True
                            break

    def insert_languages(self, language, language_level):
        inserted = False
        for table in self.cv_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if '__language__' in para.text:
                            if not inserted:
                                para.text = para.text.replace("__language__", language)
                                para.text = para.text.replace("__speaking_level__", language_level)
                                for run in para.runs:
                                    run.font.size = Pt(8)
                                inserted = True
                            #break

    def delete_tables(self):
        for table in self.cv_doc.tables:
            delete_table = False
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if '__start_date__' in para.text:
                            delete_table = True
            if delete_table:
                table_element = table._element
                table_element.getparent().remove(table_element)
                continue

    def delete_paras(self):
        for table in self.cv_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    delete_next = False
                    for para in cell.paragraphs:
                        if delete_next:
                            para_element = para._element
                            para_element.getparent().remove(para_element)
                            delete_next = False
                        if '__empty__' in para.text:
                            para_element = para._element
                            para_element.getparent().remove(para_element)
                            delete_next = True
                        if '__language__' in para.text:
                            para_element = para._element
                            para_element.getparent().remove(para_element)
                            delete_next = True
                        if '__top_row__' in para.text:
                            para_element = para._element
                            para_element.getparent().remove(para_element)
                            delete_next = True
                        if '__responsibility__' in para.text or '__achievement__' in para.text:
                            para_element = para._element
                            para_element.getparent().remove(para_element)

    def remove_extra_paras(self):
        for para in reversed(self.cv_doc.paragraphs):
            if not para.text and not has_image(para):
                para_element = para._element
                para_element.getparent().remove(para_element)

    def remove_table_paras_paras(self):
        for table in self.cv_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text and not has_image(para):
                            para_element = para._element
                            para_element.getparent().remove(para_element)

    def merge_experience(self):
        for table in self.cv_doc.tables:
            table_id = table.rows[0].cells[0].paragraphs[0].text
            if 'work experience' in table_id.lower():
                header_row = table.rows[2]
                header_cells = header_row.cells
                left_header_cell = header_cells[0]
                right_header_cell = header_cells[1]
                content_row = table.rows[3]
                content_cells = content_row.cells
                left_content_cell = content_cells[0]
                right_content_cell = content_cells[1]
                left_content_panel = left_content_cell.paragraphs
                right_content_panel = right_content_cell.paragraphs
                if not right_content_panel or all(p.text.strip() == "" for p in right_content_panel):
                    right_content_cell.text = ""
                    left_content_cell.merge(right_content_cell)
                    right_header_cell.text = ""
                    left_header_cell.merge(right_header_cell)
                if not left_content_panel or all(p.text.strip() == "" for p in left_content_panel):
                    left_content_cell.text = ""
                    left_content_cell.merge(right_content_cell)
                    left_header_cell.text = ""
                    left_header_cell.merge(right_header_cell)

                loop_elements = [left_content_cell.paragraphs, right_content_cell.paragraphs, left_header_cell.paragraphs, right_header_cell.paragraphs]
                for loop_element in loop_elements:
                    for para in loop_element:
                        if 'w:pict' in para._element.xml:
                            continue
                        if not para.text:
                            para_element = para._element
                            para_element.getparent().remove(para_element)

    def insert_education(self):
        universities = self.user_details['education']
        certifications = self.user_details['certifications']
        for table in self.cv_doc.tables:
            table_header = table.rows[0].cells[0].paragraphs[0].text
            if 'education' in table_header.lower():
                content_row = table.rows[1]
                education_section = content_row.cells[0]
                certificates_section = content_row.cells[2]
                for section in [education_section, certificates_section]:
                    for para in section.paragraphs:
                        para_element = para._element
                        para_element.getparent().remove(para_element)
                counter = 0
                for university in universities:
                    counter += 1
                    new_liner = "\n" if counter != 1 else ""
                    university_name = university["university"]
                    specialization = university["specialization"]
                    cell_paragraph = education_section.add_paragraph()
                    run = cell_paragraph.add_run(f"{new_liner}University:")
                    run.font.size = Pt(8)
                    cell_paragraph.paragraph_format.line_spacing = Pt(13.8)
                    cell_paragraph = education_section.add_paragraph()
                    run = cell_paragraph.add_run(university_name)
                    run.bold = True
                    run.font.size = Pt(8)
                    cell_paragraph.paragraph_format.line_spacing = Pt(13.8)
                    cell_paragraph = education_section.add_paragraph()
                    run = cell_paragraph.add_run("\nSpecialization:")
                    run.font.size = Pt(8)
                    cell_paragraph.paragraph_format.line_spacing = Pt(13.8)
                    cell_paragraph = education_section.add_paragraph()
                    run = cell_paragraph.add_run(specialization)
                    run.bold = True
                    run.font.size = Pt(8)
                    cell_paragraph.paragraph_format.line_spacing = Pt(13.8)
                # ADD CERFITICATES
                if certifications:
                    cell_paragraph = certificates_section.add_paragraph()
                    run = cell_paragraph.add_run(f"Certificates:")
                    run.font.size = Pt(8)
                    cell_paragraph.paragraph_format.line_spacing = Pt(13.8)
                    counter = 0
                    for certification in certifications:
                        certification_date = certification['certificate_date']
                        certification_name = certification['certificate_name']
                        counter += 1
                        new_liner = "\n" if counter != 1 else ""
                        cell_paragraph = certificates_section.add_paragraph()
                        run = cell_paragraph.add_run(f"{new_liner}{certification_date} - {certification_name}")
                        run.bold = True
                        run.font.size = Pt(8)
                        cell_paragraph.paragraph_format.line_spacing = Pt(13.8)

    def merge_education(self):
        for table in self.cv_doc.tables:
            table_header = table.rows[0].cells[0].paragraphs[0].text
            if 'education' in table_header.lower():
                content_row = table.rows[1]
                education_section = content_row.cells[0]
                certificates_section = content_row.cells[2]
                certificates_section_panel = certificates_section.paragraphs
                if not certificates_section_panel:
                    certificates_section.text = ""
                    education_section.merge(certificates_section)

    def merge_skills(self):
        for table in self.cv_doc.tables:
            table_header = table.rows[0].cells[0].paragraphs[0].text
            if 'technical skills' in table_header.lower():
                if self.entity_type == "developer":
                    header_one_row = table.rows[1]
                    content_one_row = table.rows[2]
                    header_two_row = table.rows[4]
                    content_two_row = table.rows[5]
                    # MERGE DATABASE IF EMPTY
                    database_merge_header_cell = header_one_row.cells[2]
                    database_header_cell = header_one_row.cells[3]
                    database_merge_cell = content_one_row.cells[2]
                    databases_cell = content_one_row.cells[3]
                    if not databases_cell.text:
                        database_header_cell.text = ""
                        databases_cell.text = ""
                        database_merge_cell.merge(databases_cell)
                        database_merge_header_cell.merge(database_header_cell)
                    # MERGE CLOUD IF EMPTY
                    cloud_merge_header_cell = header_two_row.cells[1]
                    cloud_header_cell = header_two_row.cells[0]
                    cloud_merge_cell = content_two_row.cells[1]
                    cloud_cell = content_two_row.cells[0]
                    if not cloud_cell.text:
                        cloud_header_cell.text = ""
                        cloud_cell.text = ""
                        cloud_merge_cell.merge(cloud_cell)
                        cloud_merge_header_cell.merge(cloud_header_cell)
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if 'w:pict' in para._element.xml:
                                continue
                            if not para.text:
                                para_element = para._element
                                para_element.getparent().remove(para_element)

    def replace_hyperlink(self, old_url, new_url, new_text):
        doc_xml = self.cv_doc.part._element
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                     "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
        for hyperlink in doc_xml.findall(".//w:hyperlink", namespace):
            rId = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if rId and rId in self.cv_doc.part.rels:
                link = self.cv_doc.part.rels[rId].target_ref
                if link != old_url:
                    continue
                nodes = [node for node in hyperlink.findall(".//w:t", namespace) if node.text]
                for node in nodes:
                    if node.text:
                        node.text = new_text
                self.cv_doc.part.rels[rId]._target = new_url

    def generate_cv(self):
        self.get_template()
        self.generate_doc_name()
        self.map_variables()
        for variable_name, variable_value in self.variables_mapping.items():
            self.replace_text_preserving_formatting(variable_name, variable_value)
        work_experiences = self.user_details['work_experience']
        for work_experience in work_experiences:
            self.insert_experience(work_experience)
        languages = self.user_details['languages']
        for language, language_level in languages.items():
            self.insert_languages(language, language_level)
        summaries = self.user_details['summaries']
        counter = 0
        for summary in summaries:
            counter += 1
            if counter > 5:
                break
            self.insert_primary_achievements(summary)
        self.insert_education()
        self.delete_tables()
        self.delete_paras()
        self.remove_extra_paras()
        self.merge_skills()
        self.merge_experience()
        self.merge_education()
        if self.entity_type == "designer" and self.user_details['skills']['portfolio']:
            portfolio_url = self.user_details['skills']['portfolio']
            portfolio_source = urlparse(portfolio_url).hostname.replace('www.', '').split('.')[0].capitalize()
            self.replace_hyperlink("https://www.behance.net/search/projects", portfolio_url, portfolio_source)
        elif self.entity_type == "designer" and not self.user_details['skills']['portfolio']:
            print("REMOVE PORTFOLIO")
        self.cv_doc.save(f"output/docx/{self.cv_doc_name}.docx")
        return f"output/docx/{self.cv_doc_name}.docx"


class DriveConverter:
    def __init__(self, drive_service, file_path, drive_cv_name, drive_folder_url, folder_name):
        self.docx_file_path = file_path
        self.drive_folder_url = drive_folder_url
        self.folder_name = folder_name
        self.drive_file_name = drive_cv_name
        self.file_name = None
        self.service = drive_service
        self.drive_docx_id = None
        self.drive_pdf_id = None
        self.pdf_file_path = None

    def upload_file_to_drive(self, file_path, mime_type, parent_id):
        file_metadata = {
            # "name": self.file_name,
            "name": self.drive_file_name,
            "mimeType": mime_type,
            "parents": [parent_id]
        }
        media = MediaFileUpload(file_path)
        create_file = self.service.files().create(
            body=file_metadata,
            media_body=media,
            fields="id",
            supportsAllDrives=True
        ).execute()
        return create_file['id']

    def export_pdf_from_drive(self):
        request = self.service.files().export_media(fileId=self.drive_docx_id, mimeType='application/pdf')
        file = request.execute()
        self.pdf_file_path = f'output/pdf/{self.file_name}.pdf'
        with open(self.pdf_file_path, 'wb') as f:
            f.write(file)

    def delete_local_copies(self):
        file_paths = [self.pdf_file_path, self.docx_file_path]
        for file_path in file_paths:
            if not file_path:
                continue
            if os.path.exists(file_path):
                os.remove(file_path)
        folder_paths = ['templates']
        for folder_path in folder_paths:
            for file_name in os.listdir(folder_path):
                file_path = os.path.join(folder_path, file_name)
                if os.path.isfile(file_path):
                    os.remove(file_path)

    def create_google_drive_folder(self):
        file_metadata = {
            "name": self.folder_name,
            "mimeType": "application/vnd.google-apps.folder",
            "parents": ["14c-WXPvPiLFGTuNYESW42rGaA8ehPLIZ"],
            "driveId": "14c-WXPvPiLFGTuNYESW42rGaA8ehPLIZ",
            "supportsAllDrives": True,
        }

        folder = self.service.files().create(
            body=file_metadata,
            fields="id",
            supportsAllDrives=True
        ).execute()
        return folder.get("id")

    def clean_google_drive_folder(self, folder_id):
        query = f"'{folder_id}' in parents and trashed=false"
        results = self.service.files().list(
            q=query,
            fields="files(id, name)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        ).execute()
        files = results.get("files", [])
        for file in files:
            try:
                self.service.files().delete(fileId=file["id"], supportsAllDrives=True).execute()
            except Exception as e:
                pass
    def convert_docx_to_pdf(self):
        # drive_scopes = ["https://www.googleapis.com/auth/drive"]
        # creds = Credentials.from_authorized_user_file("credentials/gdrive/token.json", drive_scopes)
        # creds.refresh(Request())
        # self.service = build("drive", "v3", credentials=creds)
        if not self.drive_folder_url:
            drive_folder_id = self.create_google_drive_folder()
        else:
            drive_folder_id = self.drive_folder_url.replace("https://drive.google.com/drive/folders/", "")
            self.clean_google_drive_folder(drive_folder_id)

        self.file_name = self.docx_file_path.split("/")[len(self.docx_file_path.split("/")) - 1].split(".")[0]
        self.drive_docx_id = self.upload_file_to_drive(self.docx_file_path, "application/vnd.google-apps.document", drive_folder_id)
        self.export_pdf_from_drive()
        self.drive_pdf_id = self.upload_file_to_drive(self.pdf_file_path, "application/pdf", drive_folder_id)
        self.delete_local_copies()
        return {'docx': self.drive_docx_id, "pdf": self.drive_pdf_id, "folder_id": drive_folder_id}


def cv_generator(crm_record_id):
    print("=========" * 10)
    print(f"Starting CV Generation for CRM Record ID: {crm_record_id}")
    start_time = time.time()
    crm_record_handler = CrmEntity(crm_record_id)



    try:
        drive_scopes = ["https://www.googleapis.com/auth/drive"]
        creds = Credentials.from_authorized_user_file("credentials/gdrive/token.json", drive_scopes)
        creds.refresh(Request())
        drive_service = build("drive", "v3", credentials=creds)
        print(f"\tGetting Data from Zoho CRM")
        cv_data = crm_record_handler.get_cv_data()
        cv_handler = CurriculumVitae(drive_service, cv_data['role'], cv_data)
        print(f"\tInserting Data to DOCX Template")
        docx_cv_path = cv_handler.generate_cv()
        converter = DriveConverter(drive_service, docx_cv_path, cv_data["cv_record_name"], cv_data["drive_folder_url"], cv_data['full_name'])
        print(f"\tUploading DOCX to Google Drive and Converting to PDF")
        new_file_ids = converter.convert_docx_to_pdf()
    except Exception as e:
        print(f"\tError Occured: {e}")
        print(traceback.format_exc())
        new_file_ids = {'docx': None, "pdf": None, "folder_id": None}
    print(f"\tCV Generation Completed. Results: {new_file_ids}")
    print(f"\tUploading Results to Zoho CRM")
    crm_record_handler.update_cvs(new_file_ids)
    end_time = time.time()
    elapsed_time = end_time - start_time
    print(f"Time Taken to Generate CVs: {elapsed_time:.2f} seconds")
    return {
        "result": new_file_ids,
        "duration": elapsed_time
    }


# if __name__ == '__main__':
#     results = cv_generator("1576533000404410047")
#     print(results)
